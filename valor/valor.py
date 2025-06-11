import sys
from pathlib import Path

sys.path.append(str(Path(__file__).parent.parent))

from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def criar_valores(dados, doc):
    # Estilo básico
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(10)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5

    # Título
    run = p.add_run("VALORES DECLARADOS")
    run.bold = True
    run.underline = True

    # Bloco 1: VALOR AJUSTADO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    p.add_run("VALOR AJUSTADO:\n")
    p.add_run(dados["valor_ajustado"]).bold = True

    # Bloco 2: VALOR VENAL
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    p.add_run("VALOR VENAL:\n")
    p.add_run(dados["valor_venal"]).bold = True

    # Bloco 3: ITBI
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    p.add_run("VALOR ITBI (compra e venda):\n")
    p.add_run(dados["valor_itbi"]).bold = True

    print("✅ Valores adicionados ao documento.")

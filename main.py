import sys
from pathlib import Path

sys.path.append(str(Path(__file__).parent.parent))

from utils.utils import escolher_local_salvamento
from docx import Document

from qualificacoes.solteiro import (
    criar_qualificacao_solteiro,
    pergunta_dados_solteiro,
)
from qualificacoes.casado_simples import (
    criar_qualificacao_casado_simples,
    pergunta_dados_casado_simples,
)
from qualificacoes.casado_completa import (
    criar_qualificacao_casado_completa,
    pergunta_dados_casado_completa,
)
from qualificacoes.casado_anuindo import (
    criar_qualificacao_casado_anuente,
    pergunta_dados_casado_anuente,
)
from qualificacoes.divorciado import (
    criar_qualificacao_divorciado,
    pergunta_dados_divorciado,
)
from qualificacoes.viuvo import (
    criar_qualificacao_viuvo,
    pergunta_dados_viuvo,
)

from utils.coleta_dados import coletar_dados_valores
from valor.valor import criar_valores


def main():
    doc = Document()
    assinantes = []
    while True:
        # print("\n" + "=" * 58)
        # print(" " * 20 + "MENU DE OPÇÕES" + " " * 20)
        # print("=" * 58)
        # print(f"{'ESTADOS CIVIS':<16} | {'CASADOS':<22} | {'VALORES':<15}")
        # print("-" * 58)
        # print(f"[1] Solteiro{'':<4} | [4] Casado - Simples{'':<2} | [7] Valores{'':<6}")
        # print(f"[2] Divorciado{'':<2} | [5] Casado - Completo{'':<1} | {'':<15}")
        # print(f"[3] Viúvo{'':<7} | [6] Casado - Anuente{'':<2} | {'':<15}")
        # print("=" * 58)
        # print("[0] Finalizar e salvar" + " " * 22)
        # print("=" * 58)
        # opcao = input("Escolha uma opção: ").strip()

        print("\n" + "=" * 65)
        print("✨ MENU DE OPÇÕES ✨".center(65))
        print("=" * 65)
        print(f"{'🧑 ESTADOS CIVIS':<18} | {'💑 CASADOS':<23} | {'💰 VALORES':<15}")
        print("-" * 65)
        print(f"[1] Solteiro{'':<5}🚶 | [4] Casado - Simples{'':<2}💏 | [7] Valores{'':<1} 💵")
        print(
            f"[2] Divorciado{'':<2} 💔 | [5] Casado - Completo{'':<1}👪 | {'':<15}"
        )
        print(f"[3] Viúvo{'':<7} ⚰️  | [6] Casado - Anuente{'':<1} ✍️  | {'':<15}")
        print("=" * 65)
        print("✅ [0] Finalizar e salvar ✅".center(65))
        print("=" * 65)
        opcao = input(" ➡️  Escolha uma opção: ").strip()

        if opcao == "1":
            dados = pergunta_dados_solteiro()
            criar_qualificacao_solteiro(dados, doc)
            assinantes.append(dados["nome"].upper())

        elif opcao == "2":
            dados = pergunta_dados_divorciado()
            criar_qualificacao_divorciado(dados, doc)
            assinantes.append(dados["nome"].upper())

        elif opcao == "3":
            dados = pergunta_dados_viuvo()
            criar_qualificacao_viuvo(dados, doc)
            assinantes.append(dados["nome"].upper())

        elif opcao == "4":
            dados = pergunta_dados_casado_simples()
            criar_qualificacao_casado_simples(dados, doc)
            assinantes.append(dados["nome"].upper())

        elif opcao == "5":
            dados = pergunta_dados_casado_completa()
            criar_qualificacao_casado_completa(dados, doc)
            assinantes.append(dados["nome"].upper())
            assinantes.append(dados["nome_conjuge"].upper())

        elif opcao == "6":
            dados = pergunta_dados_casado_anuente()
            criar_qualificacao_casado_anuente(dados, doc)
            assinantes.append(dados["nome"].upper())
            assinantes.append(dados["nome_conjuge"].upper())

        elif opcao == "7":
            dados = coletar_dados_valores()
            criar_valores(dados, doc)

        elif opcao == "0":
            caminho = escolher_local_salvamento("Qualificações.docx")
            if caminho:
                if assinantes:
                    p_assinantes = doc.add_paragraph()

                    run = p_assinantes.add_run("(aa) " + " // ".join(assinantes))
                    run.italic = True
            doc.save(caminho)
            break
        else:
            print("\n\n⚠️  Opção inválida. Tente novamente.")


if __name__ == "__main__":
    main()

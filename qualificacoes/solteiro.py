import sys
from pathlib import Path

sys.path.append(str(Path(__file__).parent.parent))

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from utils.coleta_dados import coletar_dados_pessoa
from utils.utils import coletar_endereco_interativo, cartorio, formatar_data


# QUALFIICAÇÃO SOLTEIRO
def criar_qualificacao_solteiro(dados, caminho):
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(10)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5

    # Tratamento
    p.add_run(f"{dados['tratamento']} ")

    # Nome
    run = p.add_run(dados["nome"])
    run.bold = True
    run.underline = True

    # Qualificação Pessoal
    p.add_run(
        f", {dados['nacionalidade']}, {dados['profissao']}, {dados['filh']} de {dados['filiacao']}, {dados['portad']} da "
    )

    run = p.add_run(f"Cédula de Identidade RG n. {dados['rg']} {dados['org_exp']}, ")
    run.bold = True

    p.add_run(f"{dados['inscrit']} no ")
    run = p.add_run(f"CPF/MF sob n. {dados['cpf']}")
    run.bold = True

    p.add_run(
        ", conforme declarado e cujas cópias dos documentos ficam arquivadas nesta Serventia, "
    )

    # Estado Civil
    run = p.add_run(f"{dados["solteir"]}")
    run.bold = True

    run = p.add_run(", maior e capaz")
    run.italic = True

    # Certidão de Nascimento
    p.add_run(f", {dados['nascid']} aos {dados['nascimento']}, conforme ")
    run = p.add_run("certidão de nascimento")
    run.underline = True

    p.add_run(
        f" que ora me é apresentada, lavrada sob n. {dados["termo_certidao"]}, "
        f"às fls. {dados['folhas']}, "
        f"do Livro {dados['livro']}, "
        f"do Oficial de Registro Civil das Pessoas Naturais {dados['cartorio']}, "
    )

    # União Estável
    p.add_run("declarando neste ato que não convive em união estável, ")

    # Endereço
    p.add_run(
        f"residente e {dados['domiciliad']} na {dados['rua']}, "
        f"n. {dados['numero']}, "
        f"{dados['bairro']}, "
        f"{dados['complemento']}"
        f"{dados['nesta-na']} cidade de {dados['cidade']}/{dados['estado']}, "
        f"CEP: {dados['cep']}."
    )

    # Salva o documento
    doc.save(caminho)
    print(f"✅ Documento salvo como: {caminho}")


def pergunta_dados_solteiro():
    dados = coletar_dados_pessoa()

    # Certidão de Nascimento ==================================================
    dados["nascimento"] = formatar_data(input("Data de nascimento: "))
    dados["termo_certidao"] = input("Termo da Certidão: ")
    dados["folhas"] = input("fls.: ")
    dados["livro"] = "A-" + input("Livro: ")
    dados["cartorio"] = cartorio()

    # Endereço ================================================================
    endereco = coletar_endereco_interativo()
    dados.update(endereco)

    return dados

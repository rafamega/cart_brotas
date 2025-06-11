import sys
from pathlib import Path

sys.path.append(str(Path(__file__).parent.parent))

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from utils.coleta_dados import coletar_dados_pessoa, regime_casamento
from utils.utils import (
    unir_outorgante_e_conjuge,
    coletar_endereco_interativo,
    cartorio,
    formatar_data,
)


# QUALIFICAÇÃO CASADO-SIMPLES =================================================
def criar_qualificacao_casado_anuente(dados, doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(10)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5

    # Tratamento
    p.add_run(f"{dados['tratamento']} ")

    # Nome
    run = p.add_run(dados["nome"])
    run.bold = True
    run.underline = True

    # Qualificação Pessoal
    p.add_run(
        f", {dados['nacionalidade']}, {dados['profissao']}, {dados['filh']} de {dados['filiacao']}, {dados['portad']} da "
    )

    run = p.add_run(f"Cédula de Identidade RG n. {dados['rg']} {dados['org_exp']}, ")
    run.bold = True

    p.add_run(f"{dados['inscrit']} no ")
    run = p.add_run(f"CPF/MF sob n. {dados['cpf']}")
    run.bold = True

    p.add_run(
        ", conforme declarado e cujas cópias dos documento ficam arquivadas nesta Serventia, "
    )

    # E sua esposa/marido
    run = p.add_run(f"e {dados['sm']}, que neste ato comparece como ANUENTE, ")
    run.italic = True

    # Tratamento Cônjuge
    p.add_run(f"{dados['tratamento_conjuge']} ")

    # Nome Cônjuge
    run = p.add_run(dados["nome_conjuge"])
    run.bold = True
    run.underline = True

    # Qualificação Pessoal Cônjuge
    p.add_run(
        f", {dados['nacionalidade_conjuge']}, {dados['profissao_conjuge']}, {dados['filh_conjuge']} de {dados['filiacao_conjuge']}, {dados['portad_conjuge']} da "
    )

    run = p.add_run(
        f"Cédula de Identidade RG n. {dados['rg_conjuge']} {dados['org_exp_conjuge']}, "
    )
    run.bold = True

    p.add_run(f"{dados['inscrit_conjuge']} no ")
    run = p.add_run(f"CPF/MF sob n. {dados['cpf_conjuge']}")
    run.bold = True

    p.add_run(
        ", conforme declarado e cujas cópias dos documentos ficam arquivadas nesta Serventia, "
    )

    # Regime do Casamento
    run = p.add_run("casados")
    run.bold = True

    p.add_run(" sob o ")

    run = p.add_run(f"Regime da {dados["regime_casamento"]}")
    run.bold = True

    p.add_run(
        f", aos {dados['data_casamento']}, na vigência da Lei 6.515/77, conforme "
    )
    run = p.add_run("certidão de casamento")
    run.underline = True

    p.add_run(
        f" que ora me é apresentada, lavrada sob n. {dados['termo_certidao']}, às fls. {dados['folhas']}, "
        f"do Livro {dados['livro']}, do Oficial de Registro Civil das Pessoas Naturais {dados['cartorio']}, "
        f"{dados['pacto']}"
        f"residente e {dados['domiciliad']} na {dados['rua']}, n. {dados['numero']}, "
        f"{dados['bairro']}, {dados['complemento']}{dados['nesta-na']} cidade de "
        f"{dados['cidade']}/{dados['estado']}, CEP: {dados['cep']}."
    )

    # Salva o documento
    print(f"✅ Qualificação de {dados['nome']} adicionada ao documento.")


# ENTRADA INTERATIVA ==========================================================
def pergunta_dados_casado_anuente():
    outorgante = coletar_dados_pessoa("outorgante")
    conjuge = coletar_dados_pessoa("cônjuge")
    dados = unir_outorgante_e_conjuge(outorgante, conjuge)

    # Regime e certidão
    txt_pacto = (
        f"e seu respectivo pacto antenupcial lavrado neste Tabelionato no Livro ____, "
        f"às fls. ___, aos _______, e que foi registrado sob n. ______, livro 3, "
        f"na Oficial de Registro de Imóveis desta Comarca de Brotas/SP, "
    )

    dados["regime_casamento"] = regime_casamento()
    dados["pacto"] = (
        txt_pacto if dados["regime_casamento"] != "Comunhão Parcial de Bens" else ""
    )
    dados["data_casamento"] = formatar_data(input("Data do casamento: "))
    dados["termo_certidao"] = input("Termo: ")
    dados["folhas"] = input("Fls.: ")
    dados["livro"] = "B-" + input("Livro: ")
    dados["cartorio"] = cartorio()

    endereco = coletar_endereco_interativo()
    dados.update(endereco)

    return dados
